import os
import requests
import json
import io
from flask import Flask, request, jsonify, render_template, send_file
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 加载环境变量
load_dotenv()

app = Flask(__name__)

# Dify API相关配置
DIFY_API_KEY = os.getenv("DIFY_API_KEY")
DIFY_API_BASE_URL = os.getenv("DIFY_API_BASE_URL", "http://localhost/v1")

HEADERS = {
    "Authorization": f"Bearer {DIFY_API_KEY}",
    "Content-Type": "application/json"
}

# 项目列表
SUBJECTS = [
    "二十大精神",
    "铁路安全生产规章",
    "铁路专业技术知识",
    "应急处置流程",
    "客运服务规范",
    "设备操作维护",
    "安全事故案例分析",
    "铁路发展战略"
]

# 难度级别
DIFFICULTY_LEVELS = [
    "简单", "中等", "困难"
]

# 题型列表
QUESTION_TYPES = [
    "选择题", "填空题", "判断题", "简答题", "计算题", "阅读理解题"
]

@app.route("/")
def index():
    """渲染首页"""
    return render_template("index.html", 
                           subjects=SUBJECTS, 
                           difficulty_levels=DIFFICULTY_LEVELS,
                           question_types=QUESTION_TYPES)

@app.route("/api/generate-questions", methods=["POST"])
def generate_questions():
    """生成题目的API接口"""
    try:
        data = request.json
        
        # 获取请求参数
        subject = data.get("subject", "二十大精神")
        difficulty = data.get("difficulty", "中等")
        question_type = data.get("question_type", "选择题")
        grade = data.get("grade", "初一")
        count = min(int(data.get("count", 5)), 10)  # 限制最多10题
        
        # 构建提问内容
        prompt = f"请生成{count}道{grade}{subject}{difficulty}难度的{question_type}，包含答案和解析。格式要求：每道题的题目、选项(如果是选择题)、答案和解析要分开展示，便于阅读。"
        
        # 调用Dify API
        response = requests.post(
            f"{DIFY_API_BASE_URL}/chat-messages",
            headers=HEADERS,
            json={
                "inputs": {},
                "query": prompt,
                "response_mode": "blocking",  # 阻塞模式，等待完整回复
                "user": "question-generator-app"  # 用户标识
            }
        )
        
        response.raise_for_status()
        result = response.json()
        
        # 提取回答内容
        answer_content = result.get("answer", "无法生成题目，请稍后再试")
        
        # 添加当前时间戳到元数据
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        return jsonify({
            "success": True,
            "questions": answer_content,
            "metadata": {
                "subject": subject,
                "difficulty": difficulty,
                "grade": grade,
                "question_type": question_type,
                "count": count,
                "timestamp": timestamp
            }
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/conversations", methods=["GET"])
def get_conversations():
    """获取历史会话列表"""
    try:
        response = requests.get(
            f"{DIFY_API_BASE_URL}/conversations",
            headers=HEADERS,
            params={"user": "question-generator-app", "limit": 20}
        )
        
        response.raise_for_status()
        result = response.json()
        
        return jsonify({
            "success": True,
            "conversations": result.get("data", [])
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/conversation/<conversation_id>", methods=["GET"])
def get_conversation_messages(conversation_id):
    """获取特定会话的消息记录"""
    try:
        response = requests.get(
            f"{DIFY_API_BASE_URL}/messages",
            headers=HEADERS,
            params={
                "user": "question-generator-app", 
                "conversation_id": conversation_id,
                "limit": 100
            }
        )
        
        response.raise_for_status()
        result = response.json()
        
        return jsonify({
            "success": True,
            "messages": result.get("data", [])
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/history")
def history():
    """渲染历史记录页面"""
    return render_template("history.html")

@app.route("/about")
def about():
    """关于页面"""
    return render_template("about.html")

@app.route("/api/download-word", methods=["POST"])
def download_word():
    """将生成的题目转换为Word文档并下载"""
    try:
        data = request.json
        questions_text = data.get("questions", "")
        metadata = data.get("metadata", {})
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档标题
        title = f"{metadata.get('grade', '')} {metadata.get('subject', '')} {metadata.get('question_type', '')}习题"
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元数据
        metadata_para = doc.add_paragraph()
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_text = f"难度: {metadata.get('difficulty', '')} | 题目数量: {metadata.get('count', '')} | 生成时间: {metadata.get('timestamp', '')}"
        metadata_para.add_run(metadata_text).italic = True
        
        # 添加分隔线
        doc.add_paragraph("=" * 50)
        
        # 添加题目内容
        doc.add_paragraph(questions_text)
        
        # 保存到内存流
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 生成文件名
        filename = f"{metadata.get('grade', '中等')}{metadata.get('subject', '二十大精神')}{metadata.get('question_type', '题目')}.docx"
        
        # 返回文件
        return send_file(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
